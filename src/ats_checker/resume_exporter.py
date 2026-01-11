"""Resume export functionality for multiple formats."""

import io
import logging
from pathlib import Path

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)


class ResumeExporter:
    """Export resumes to various formats."""

    def __init__(self, content: str):
        """
        Initialize exporter with resume content.

        Args:
            content: Resume text content (plain text or markdown)
        """
        self.content = content

    def to_markdown(self) -> str:
        """
        Convert resume to clean Markdown format.

        Returns:
            Markdown-formatted resume
        """
        lines = self.content.split("\n")
        md_lines: list[str] = []
        in_list = False

        for line in lines:
            stripped = line.strip()

            # Skip empty lines but preserve spacing
            if not stripped:
                if in_list:
                    in_list = False
                md_lines.append("")
                continue

            # Detect section headers (ALL CAPS or ending with :)
            if stripped.isupper() and len(stripped) > 2:
                md_lines.append(f"\n## {stripped.title()}\n")
            elif stripped.endswith(":") and len(stripped.split()) <= 4:
                md_lines.append(f"\n### {stripped[:-1]}\n")
            # Detect bullet points
            elif stripped.startswith(("-", "•", "*", "–")):
                md_lines.append(f"- {stripped[1:].strip()}")
                in_list = True
            # Detect numbered items
            elif stripped[0].isdigit() and "." in stripped[:3]:
                md_lines.append(stripped)
                in_list = True
            else:
                md_lines.append(stripped)

        return "\n".join(md_lines)

    def to_html(self, include_style: bool = True) -> str:
        """
        Convert resume to HTML.

        Args:
            include_style: Whether to include CSS styling

        Returns:
            HTML string
        """
        md_content = self.to_markdown()
        html_content = markdown.markdown(md_content, extensions=["tables", "fenced_code", "nl2br"])

        if include_style:
            style = """
            <style>
                body {
                    font-family: 'Helvetica Neue', Arial, sans-serif;
                    max-width: 800px;
                    margin: 40px auto;
                    padding: 20px;
                    line-height: 1.6;
                    color: #333;
                }
                h1 {
                    color: #2c3e50;
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 10px;
                }
                h2 {
                    color: #2c3e50;
                    margin-top: 25px;
                    border-bottom: 1px solid #bdc3c7;
                    padding-bottom: 5px;
                }
                h3 {
                    color: #34495e;
                    margin-top: 15px;
                }
                ul {
                    padding-left: 20px;
                }
                li {
                    margin-bottom: 5px;
                }
                p {
                    margin: 10px 0;
                }
                @media print {
                    body { margin: 0; padding: 15px; }
                }
            </style>
            """
            return f"<!DOCTYPE html><html><head><meta charset='utf-8'>{style}</head><body>{html_content}</body></html>"

        return html_content

    def to_pdf(self) -> bytes:
        """
        Convert resume to PDF using WeasyPrint.

        Returns:
            PDF as bytes
        """
        try:
            from weasyprint import HTML

            html_content = self.to_html(include_style=True)
            pdf_buffer = io.BytesIO()
            HTML(string=html_content).write_pdf(pdf_buffer)
            pdf_buffer.seek(0)
            return pdf_buffer.read()
        except ImportError:
            logger.error("WeasyPrint not installed. Cannot generate PDF.")
            raise RuntimeError(
                "PDF export requires WeasyPrint. Install with: pip install weasyprint"
            )
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise RuntimeError(f"Failed to generate PDF: {e}") from e

    def to_docx(self) -> bytes:
        """
        Convert resume to DOCX format.

        Returns:
            DOCX as bytes
        """
        doc = Document()

        # Set up styles
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        lines = self.content.split("\n")

        for line in lines:
            stripped = line.strip()

            if not stripped:
                # Add empty paragraph for spacing
                doc.add_paragraph()
                continue

            # Detect section headers
            if stripped.isupper() and len(stripped) > 2:
                heading = doc.add_heading(stripped.title(), level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                stripped.lower()
            elif stripped.endswith(":") and len(stripped.split()) <= 4:
                doc.add_heading(stripped[:-1], level=2)
            # Bullet points
            elif stripped.startswith(("-", "•", "*", "–")):
                para = doc.add_paragraph(stripped[1:].strip(), style="List Bullet")
            # Regular paragraph
            else:
                para = doc.add_paragraph(stripped)

                # Make first line (usually name) bold and larger
                if lines.index(line) == 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(16)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def save(self, path: str | Path, format: str = "pdf") -> Path:
        """
        Save resume to file.

        Args:
            path: Output file path (without extension)
            format: Export format ('markdown', 'pdf', 'docx', 'html')

        Returns:
            Path to saved file
        """
        path = Path(path)

        if format == "markdown" or format == "md":
            output_path = path.with_suffix(".md")
            output_path.write_text(self.to_markdown(), encoding="utf-8")
        elif format == "pdf":
            output_path = path.with_suffix(".pdf")
            output_path.write_bytes(self.to_pdf())
        elif format == "docx":
            output_path = path.with_suffix(".docx")
            output_path.write_bytes(self.to_docx())
        elif format == "html":
            output_path = path.with_suffix(".html")
            output_path.write_text(self.to_html(), encoding="utf-8")
        else:
            raise ValueError(f"Unsupported format: {format}")

        logger.info(f"Resume saved to {output_path}")
        return output_path


def export_resume(content: str, format: str = "pdf") -> bytes | str:
    """
    Convenience function to export resume content.

    Args:
        content: Resume text content
        format: Export format

    Returns:
        Exported content (bytes for pdf/docx, str for md/html)
    """
    exporter = ResumeExporter(content)

    if format == "markdown" or format == "md":
        return exporter.to_markdown()
    elif format == "pdf":
        return exporter.to_pdf()
    elif format == "docx":
        return exporter.to_docx()
    elif format == "html":
        return exporter.to_html()
    else:
        raise ValueError(f"Unsupported format: {format}")
