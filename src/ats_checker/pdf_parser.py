"""PDF parsing utilities for extracting text from resume and job description PDFs."""

import io
import logging
import re
from pathlib import Path
from typing import BinaryIO

import pdfplumber

logger = logging.getLogger(__name__)


def _extract_text_with_char_spacing(page: "pdfplumber.page.Page") -> str:
    """
    Extract text using character-level analysis to properly handle spacing.

    This method extracts individual characters and reconstructs words based on
    actual spacing between characters, which works better for PDFs with tight
    kerning or unusual font metrics.
    """
    chars = page.chars
    if not chars:
        return ""

    # Sort characters by vertical position (top) then horizontal (x0)
    chars = sorted(chars, key=lambda c: (round(c["top"], 1), c["x0"]))

    lines: list[list[dict]] = []
    current_line: list[dict] = []
    last_top = None
    y_threshold = 5  # Characters within this vertical distance are on same line

    for char in chars:
        if last_top is None or abs(char["top"] - last_top) <= y_threshold:
            current_line.append(char)
        else:
            if current_line:
                lines.append(current_line)
            current_line = [char]
        last_top = char["top"]

    if current_line:
        lines.append(current_line)

    # Analyze gaps across all lines to find optimal space threshold
    all_gaps: list[float] = []
    for line_chars in lines:
        sorted_chars = sorted(line_chars, key=lambda c: c["x0"])
        for i in range(1, len(sorted_chars)):
            gap = sorted_chars[i]["x0"] - sorted_chars[i - 1]["x1"]
            if gap > 0:
                all_gaps.append(gap)

    # Find a good threshold: typically there's a cluster of small gaps (within words)
    # and larger gaps (between words). Use a percentile-based approach.
    if all_gaps:
        all_gaps.sort()
        # Use 40th percentile as threshold - more aggressive separation
        # Most PDF tight-kerned text has gaps clustered near 0, spaces are larger
        threshold_idx = int(len(all_gaps) * 0.40)
        space_threshold = all_gaps[threshold_idx] if threshold_idx < len(all_gaps) else 1.0
        # Minimum threshold to avoid splitting within letters
        space_threshold = max(space_threshold, 0.5)
    else:
        space_threshold = 1.5

    # Process each line
    text_lines: list[str] = []
    for line_chars in lines:
        # Sort by horizontal position
        line_chars = sorted(line_chars, key=lambda c: c["x0"])

        line_text = ""
        last_x1 = None

        for char in line_chars:
            char_text = char.get("text", "")
            if not char_text:
                continue

            if last_x1 is not None:
                # Calculate gap between characters
                gap = char["x0"] - last_x1

                # Use dynamic threshold
                if gap > space_threshold:
                    line_text += " "

            line_text += char_text
            last_x1 = char["x1"]

        if line_text.strip():
            text_lines.append(line_text.strip())

    return "\n".join(text_lines)


def extract_text_from_pdf(
    source: str | Path | BinaryIO | bytes,
    extract_tables: bool = True,
    pages: list[int] | None = None,
) -> str:
    """
    Extract text content from a PDF file.

    Args:
        source: Path to PDF file, file-like object, or bytes
        extract_tables: Whether to extract table content as well
        pages: Specific page numbers to extract (0-indexed), None for all pages

    Returns:
        Extracted text content as a string

    Raises:
        ValueError: If the PDF cannot be read or is empty
    """
    text_parts: list[str] = []

    try:
        # Handle different source types
        if isinstance(source, bytes):
            pdf_file: BinaryIO = io.BytesIO(source)
        elif isinstance(source, (str, Path)):
            pdf_file = open(source, "rb")  # noqa: SIM115
        else:
            pdf_file = source

        with pdfplumber.open(pdf_file) as pdf:
            if not pdf.pages:
                raise ValueError("PDF file contains no pages")

            pages_to_process = (
                [pdf.pages[i] for i in pages if i < len(pdf.pages)] if pages else pdf.pages
            )

            for page in pages_to_process:
                # Use character-level extraction for better spacing
                page_text = _extract_text_with_char_spacing(page)

                if page_text.strip():
                    # Post-process to fix remaining spacing issues
                    page_text = _fix_text_spacing(page_text)
                    text_parts.append(page_text)

                # Extract tables if requested
                if extract_tables:
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = _format_table(table)
                        if table_text.strip():
                            text_parts.append(table_text)

        # Close file if we opened it
        if isinstance(source, (str, Path)):
            pdf_file.close()

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}") from e

    if not text_parts:
        raise ValueError("No text content found in PDF")

    return "\n\n".join(text_parts)


def _fix_text_spacing(text: str) -> str:
    """
    Fix common text spacing issues from PDF extraction.

    Args:
        text: Raw extracted text

    Returns:
        Text with improved spacing
    """
    # Fix words that are stuck together (camelCase-like patterns in normal text)
    # e.g., "experienceWorking" -> "experience Working"
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)

    # Fix missing space after periods (but not for abbreviations like "U.S.")
    text = re.sub(r"\.([A-Z][a-z])", r". \1", text)

    # Fix missing space after commas
    text = re.sub(r",([A-Za-z])", r", \1", text)

    # Fix multiple spaces
    text = re.sub(r" +", " ", text)

    # Fix missing space between number and text (e.g., "5years" -> "5 years")
    text = re.sub(r"(\d)([a-zA-Z])", r"\1 \2", text)
    text = re.sub(r"([a-zA-Z])(\d)", r"\1 \2", text)

    # Clean up extra whitespace
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if line:
            lines.append(line)

    return "\n".join(lines)


def _format_table(table: list[list[str | None]]) -> str:
    """Format a table as readable text."""
    if not table:
        return ""

    lines: list[str] = []
    for row in table:
        # Filter out None values and join with tabs
        cells = [str(cell) if cell else "" for cell in row]
        if any(cell.strip() for cell in cells):
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_text_from_docx(source: str | Path | BinaryIO | bytes) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        source: Path to DOCX file, file-like object, or bytes

    Returns:
        Extracted text content as a string
    """
    from docx import Document

    try:
        if isinstance(source, bytes):
            doc = Document(io.BytesIO(source))
        elif isinstance(source, (str, Path)):
            doc = Document(source)
        else:
            doc = Document(source)

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            raise ValueError("No text content found in DOCX")

        return "\n\n".join(paragraphs)

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}") from e


def extract_text_from_file(
    source: str | Path | BinaryIO | bytes,
    file_name: str | None = None,
) -> str:
    """
    Extract text from a file, auto-detecting the format.

    Args:
        source: File path, file-like object, or bytes
        file_name: Original file name (used for format detection when source is bytes)

    Returns:
        Extracted text content
    """
    # Determine file extension
    if isinstance(source, (str, Path)):
        ext = Path(source).suffix.lower()
    elif file_name:
        ext = Path(file_name).suffix.lower()
    else:
        ext = ""

    if ext == ".pdf":
        return extract_text_from_pdf(source)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(source)
    elif ext in (".txt", ".md", ""):
        # Plain text
        if isinstance(source, bytes):
            return source.decode("utf-8")
        elif isinstance(source, (str, Path)):
            return Path(source).read_text(encoding="utf-8")
        else:
            content = source.read()
            if isinstance(content, bytes):
                return content.decode("utf-8")
            return content
    else:
        raise ValueError(f"Unsupported file format: {ext}")
